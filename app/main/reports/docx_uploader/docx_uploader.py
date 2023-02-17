import re
from functools import reduce
from typing import List

import docx

from app.main.reports.pdf_document.pdf_document_manager import PdfDocumentManager

from .core_properties import CoreProperties
from .inline_shape import InlineShape
from .paragraph import Paragraph
from .table import Table, Cell
from .style import Style
from ..pdf_document.pdf_document_manager import PdfDocumentManager
from ...checks.report_checks.style_check_settings import StyleCheckSettings



class DocxUploader:
    def __init__(self):
        self.inline_shapes = []
        self.core_properties = None
        self.chapters = []
        self.paragraphs = []
        self.tables = []
        self.file = None
        self.styled_paragraphs = None
        self.special_paragraph_indices = {}
        self.pdf_file = None
        self.styles: List[Style] = []
        self.count = 0

    def upload(self, file):
        self.file = docx.Document(file)
        self.pdf_file = PdfDocumentManager(file)

    def parse(self):
        self.core_properties = CoreProperties(self.file)
        for i in range(len(self.file.inline_shapes)):
            self.inline_shapes.append(InlineShape(self.file.inline_shapes[i]))
        self.paragraphs = self.__make_paragraphs(self.file.paragraphs)
        self.parse_effective_styles()
        self.tables = self.__make_table(self.file.tables)

    def __make_paragraphs(self, paragraphs):
        tmp_paragraphs = []
        for i in range(len(paragraphs)):
            if len(paragraphs[i].text.strip()):
                tmp_paragraphs.append(Paragraph(paragraphs[i]))
        return tmp_paragraphs

    def make_chapters(self, work_type):
        headers = []
        tmp_chapters = []

        if work_type == 'VKR':
            # find first pages
            headers = [
                {"name": "Титульный лист", "marker": False, "key": "санкт-петербургский государственный", "page": 0},
                {"name": "Задание на выпускную квалификационную работу", "marker": False, "key": "задание", "page": 0},
                {"name": "Календарный план", "marker": False, "key": "календарный план", "page": 0},
                {"name": "Реферат", "marker": False, "key": "реферат", "page": 0},
                {"name": "Abstract", "marker": False, "key": "abstract", "page": 0},
                {"name": "Cодержание", "marker": False, "key": "содержание", "page": 0}]
            page = 1
            elem = 0
            while elem < len(headers) and page < (2 * len(headers)):
                page_text = (self.pdf_file.get_text_on_page()[page].split("\n")[0]).lower()
                if page_text.find(headers[elem]["key"]) >= 0:
                    headers[elem]["marker"] = True
                    headers[elem]["page"] = page
                    elem += 1
                else:
                    page_text = (self.pdf_file.get_text_on_page()[page + 1].split("\n")[0]).lower()
                    if page_text.find(headers[elem]["key"]) >= 0:
                        headers[elem]["marker"] = True
                        headers[elem]["page"] = page + 1
                        elem += 1
                        page += 1
                page += 1

            # find headers
            header_ind = -1
            par_num = 0
            head_par_ind = -1
            for par_ind in range(len(self.styled_paragraphs)):
                head_par_ind += 1
                style_name = self.paragraphs[par_ind].paragraph_style_name.lower()
                if style_name.find("heading") >= 0:
                    print(self.paragraphs[par_ind].paragraph_style_name)
                    print(self.paragraphs[par_ind].paragraph_text)
                    print(self.styled_paragraphs[par_ind]["text"])
                    header_ind += 1
                    par_num = 0
                    tmp_chapters.append({"style": style_name, "text": self.styled_paragraphs[par_ind]["text"], "styled_text": self.styled_paragraphs[par_ind], "number": head_par_ind, "child": []})
                elif header_ind >= 0:
                    par_num += 1
                    tmp_chapters[header_ind]["child"].append({"style": style_name, "text": self.styled_paragraphs[par_ind]["text"], "styled_text": self.styled_paragraphs[par_ind], "number": head_par_ind})
        self.chapters = tmp_chapters
        return headers

    def __make_table(self, tables):
        for i in range(len(tables)):
            table = []
            for j in range(len(tables[i].rows)):
                row = []
                for k in range(len(tables[i].rows[j].cells)):
                    tmp_paragraphs = self.__make_paragraphs(tables[i].rows[j].cells[k].paragraphs)
                    row.append(Cell(tables[i].rows[j].cells[k], tmp_paragraphs))
                table.append(row)
            self.tables.append(Table(tables[i], table))
        return tables

    def build_vkr_hierarchy(self):
        indices = self.get_paragraph_indices_by_style(self.styles)
        tagged_indices = [{"index": 0, "level": 0}, {"index": len(self.styled_paragraphs), "level": 0}]
        for j in range(len(indices)):
            tagged_indices.extend(list(map(lambda index: {"index": index, "level": j + 1,
                                                          "text": self.styled_paragraphs[index]["text"]}, indices[j])))
        tagged_indices.sort(key=lambda dct: dct["index"])
        return tagged_indices

    # Parses styles once; subsequent calls have no effect, since the file itself shouldn't change
    def parse_effective_styles(self):
        if self.styled_paragraphs is not None:
            return
        self.styled_paragraphs = []
        for par in filter(lambda p: len(p.text.strip()) > 0, self.file.paragraphs):
            paragraph = {"text": par.text, "runs": []}
            for run in filter(lambda r: len(r.text.strip()) > 0, par.runs):
                paragraph["runs"].append({"text": run.text, "style": Style(run, par)})
            self.styled_paragraphs.append(paragraph)

    def unify_multiline_entities(self, first_line_regex_str):
        pattern = re.compile(first_line_regex_str)
        pars_to_delete = []
        skip_flag = False
        for i in range(len(self.styled_paragraphs)-1):
            if skip_flag:
                skip_flag = False
                continue
            par = self.styled_paragraphs[i]
            next_par = self.styled_paragraphs[i+1]
            if pattern.match(par["text"]):
                skip_flag = True
                par["text"] += ("\n" + next_par["text"])
                par["runs"].extend(next_par["runs"])
                pars_to_delete.append(next_par)
                continue
        for par in pars_to_delete:
            self.styled_paragraphs.remove(par)

    def get_paragraph_indices_by_style(self, style_list):
        result = []
        for template_style in style_list:
            matched_pars = []
            for i in range(len(self.styled_paragraphs)):
                par = self.styled_paragraphs[i]
                if reduce(lambda prev, run: prev and run["style"].matches(template_style), par["runs"], True):
                    matched_pars.append(i)
            result.append(matched_pars)
        return result

    def page_counter(self):
        if not self.count:
            for k, v in self.pdf_file.text_on_page.items():
                if re.search('приложение [а-я][\n .]', v.lower()):
                    break
                self.count += 1
        return self.count


    def upload_from_cli(self, file):
        self.upload(file=file)

    def print_info(self):
        print(self.core_properties.to_string())
        for i in range(len(self.paragraphs)):
            print(self.paragraphs[i].to_string())

    def __str__(self):
        return self.core_properties.to_string() + '\n' + '\n'.join([self.paragraphs[i].to_string() for i in range(len(self.paragraphs))])


def main(args):
    file = args.file
    uploader = DocxUploader()
    uploader.upload_from_cli(file=file)
    uploader.parse()
    uploader.print_info()
    uploader.parse_effective_styles()
